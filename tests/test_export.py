import requests
from docx import Document

# Login and export
r_login = requests.post('http://localhost:8000/auth/login', json={'email': 'demo@example.com', 'password': 'demo123'})
token = r_login.json()['access_token']
r_export = requests.post('http://localhost:8000/suggestions/runs/88d6a06f-6179-4979-81eb-b2d573b6c97a/export-with-accepted', 
                         headers={'Authorization': f'Bearer {token}'})

print(f'Export status: {r_export.status_code}')
print(f'Export size: {len(r_export.content)} bytes')

if r_export.status_code == 200:
    # Save file
    with open('/tmp/exported.docx', 'wb') as f:
        f.write(r_export.content)
    
    # Read and display content
    doc = Document('/tmp/exported.docx')
    print('\n=== DOCUMENTO EXPORTADO CON CORRECCIONES ACEPTADAS ===\n')
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            print(f'{i}. {para.text}')
    
    print('\n=== VERIFICACIÓN ===')
    print('La corrección RECHAZADA (halla->haya) NO debe aparecer')
    print('Las 14 correcciones ACEPTADAS deben estar aplicadas')
    
    # Show specific checks
    full_text = ' '.join(p.text for p in doc.paragraphs)
    print(f'\nContiene "halla" (rechazada, debe estar): {"halla" in full_text}')
    print(f'Contiene "haya" (no debe estar): {"haya" in full_text}')
    print(f'Contiene "vaca" (aceptada): {"vaca" in full_text}')
    print(f'Contiene "ha hecho" (aceptada): {"ha hecho" in full_text}')
else:
    print(f'Error: {r_export.text}')

