from __future__ import annotations

import json
import logging
from collections.abc import Iterable, Sequence
from dataclasses import dataclass
from pathlib import Path

from .docx_utils import read_paragraphs, write_docx_preserving_runs, write_paragraphs
from .model import BaseCorrector, CorrectionSpec
from .text_utils import (
    Token,
    apply_token_corrections,
    build_context,
    build_sentence_context,
    detokenize,
    split_tokens_by_char_budget,
    split_tokens_in_chunks,
    tokenize,
)

logger = logging.getLogger(__name__)
try:  # optional rich formatting for DOCX log
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None  # type: ignore


@dataclass
class LogEntry:
    token_id: int
    line: int  # 1-based paragraph index
    original: str
    corrected: str
    reason: str
    context: str
    chunk_index: int
    sentence: str


def paragraphs_to_text(paragraphs: Sequence[str]) -> str:
    return "\n".join(paragraphs)


def text_to_paragraphs(text: str) -> list[str]:
    return text.split("\n")


def process_paragraphs(
    paragraphs: Sequence[str],
    corrector: BaseCorrector,
    *,
    chunk_words: int = 0,
    overlap_words: int = 0,
) -> tuple[list[str], list[LogEntry]]:
    # Tokenize full document text to create stable global token ids
    full_text = paragraphs_to_text(paragraphs)
    tokens = tokenize(full_text)

    # Compute chunks as ranges of token indices
    # Compute chunks as ranges of token indices
    if chunk_words and chunk_words > 0:
        ranges = split_tokens_in_chunks(tokens, max_words=chunk_words, overlap_words=overlap_words)
    else:
        # Auto-chunk by approximate character budget using ~70% of 128k tokens context
        CONTEXT_TOKENS = 128_000
        CHAR_PER_TOKEN_EST = 4
        FRACTION = 0.7
        char_budget = int(CONTEXT_TOKENS * CHAR_PER_TOKEN_EST * FRACTION)
        overlap_chars = int(char_budget * 0.03)
        ranges = split_tokens_by_char_budget(tokens, char_budget=char_budget, overlap_chars=overlap_chars)

    applied_global: dict[int, CorrectionSpec] = {}
    log_entries: list[LogEntry] = []
    total_chunks = len(ranges)
    logger.info(f"Procesando documento en {total_chunks} chunk(s)...")
    for chunk_idx, (start, end) in enumerate(ranges):
        logger.info(f"📄 Procesando chunk {chunk_idx + 1}/{total_chunks} (tokens {start}-{end})...")
        # Local ids start from 0; map back to global by +start
        local_tokens = [Token(i - start, t.text, t.start, t.end, t.kind, t.line) for i, t in enumerate(tokens[start:end], start=start)]
        logger.info(f"🔄 Enviando chunk {chunk_idx + 1}/{total_chunks} al corrector...")
        corrections = corrector.correct_tokens(local_tokens)
        logger.info(f"✅ Chunk {chunk_idx + 1}/{total_chunks}: {len(corrections)} correcciones encontradas")
        for c in corrections:
            global_id = start + c.token_id
            if 0 <= global_id < len(tokens):
                if global_id in applied_global:
                    # already applied due to overlap; skip duplicates
                    continue
                tok = tokens[global_id]

                # Skip if replacement is identical to original (false positive)
                if tok.text == c.replacement:
                    logger.debug(f"Skipping false positive: '{tok.text}' == '{c.replacement}'")
                    continue

                # Skip if original is whitespace/punctuation but replacement is a word
                # (Gemini sometimes gets token IDs wrong)
                if tok.text.strip() == "" and c.replacement.strip() != "":
                    logger.warning(f"Skipping invalid correction: whitespace token '{repr(tok.text)}' -> '{c.replacement}'")
                    continue

                if tok.kind in ("space", "newline", "punct") and c.replacement not in (" ", "\n", tok.text):
                    logger.warning(f"Skipping suspicious correction: {tok.kind} token '{tok.text}' -> '{c.replacement}'")
                    continue

                # Detectar si es una eliminación (replacement vacío o muy diferente)
                corrected_text = c.replacement
                reason_text = c.reason

                # Si el replacement parece ser el siguiente token, probablemente es una eliminación
                if global_id + 1 < len(tokens) and c.replacement.strip() == tokens[global_id + 1].text.strip():
                    corrected_text = ""
                    reason_text = f"[ELIMINACIÓN] {c.reason}"
                elif c.replacement.strip() == "":
                    reason_text = f"[ELIMINACIÓN] {c.reason}"

                entry = LogEntry(
                    token_id=global_id,
                    line=tok.line,
                    original=tok.text,
                    corrected=corrected_text,
                    reason=reason_text,
                    context=build_context(tokens, global_id, radius=3),
                    chunk_index=chunk_idx,
                    sentence=build_sentence_context(tokens, global_id),
                )
                log_entries.append(entry)
                applied_global[global_id] = c

    # Apply all corrections to the global token list
    if applied_global:
        ordered = [
            type("Corr", (), {"token_id": k, "replacement": v.replacement, "reason": v.reason, "original": v.original})
            for k, v in applied_global.items()
        ]
        tokens = apply_token_corrections(tokens, ordered)

    corrected_text = detokenize(tokens)
    corrected_paragraphs = text_to_paragraphs(corrected_text)
    return corrected_paragraphs, log_entries


def process_document(
    input_path: str,
    output_path: str,
    log_path: str,
    corrector: BaseCorrector,
    *,
    chunk_words: int = 0,
    overlap_words: int = 0,
    preserve_format: bool = True,
    log_docx_path: str | None = None,
    enable_docx_log: bool = True,
) -> None:
    paragraphs = read_paragraphs(input_path)
    corrected_paragraphs, log_entries = process_paragraphs(
        paragraphs, corrector, chunk_words=chunk_words, overlap_words=overlap_words
    )
    # Preserve formatting for DOCX outputs by rewriting document.xml text only
    if (
        preserve_format
        and output_path.lower().endswith('.docx')
        and input_path.lower().endswith('.docx')
    ):
        write_docx_preserving_runs(input_path, corrected_paragraphs, output_path)
    else:
        write_paragraphs(corrected_paragraphs, output_path)
    _write_log_jsonl(log_path, log_entries)
    if enable_docx_log:
        if log_docx_path:
            docx_path = log_docx_path
        else:
            # Si log_path está en outputs/, poner el DOCX también ahí
            log_parent = Path(log_path).parent
            if log_parent.name == "outputs" or str(log_parent).endswith("outputs"):
                docx_path = str(log_parent / f"{Path(input_path).stem}.corrections.docx")
            else:
                # Fallback: mismo directorio que el log JSONL
                docx_path = str(Path(log_path).with_suffix(".docx"))
        _write_log_docx(docx_path, log_entries, source_filename=Path(input_path).name)


def _write_log_jsonl(path: str, entries: Iterable[LogEntry]) -> None:
    p = Path(path)
    with p.open("w", encoding="utf-8") as f:
        for e in entries:
            f.write(
                json.dumps(
                    {
                        "token_id": e.token_id,
                        "line": e.line,
                        "original": e.original,
                        "corrected": e.corrected,
                        "reason": e.reason,
                        "context": e.context,
                        "chunk_index": e.chunk_index,
                        "sentence": e.sentence,
                    },
                    ensure_ascii=False,
                )
                + "\n"
            )


def _safe_get(obj, name: str, default=None):
    if hasattr(obj, name):
        return getattr(obj, name)
    if isinstance(obj, dict):
        return obj.get(name, default)
    return default


def _write_log_docx(path: str, entries: Iterable[LogEntry], *, source_filename: str | None = None) -> None:
    entries_list = list(entries)
    if Document is not None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor

        doc = Document()  # type: ignore

        # Título
        title = doc.add_heading("Informe de Correcciones", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Información del documento
        if source_filename:
            p = doc.add_paragraph()
            p.add_run("Documento: ").bold = True
            p.add_run(source_filename)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.add_run("Total de correcciones: ").bold = True
        p.add_run(str(len(entries_list)))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Espacio

        # Crear tabla con estilo
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        # Encabezados
        hdr_cells = table.rows[0].cells
        headers = ["#", "Original → Corregido", "Motivo", "Contexto", "Línea"]

        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            # Formatear encabezado
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Color de fondo del encabezado
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '4472C4')
            hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)

        # Ajustar anchos de columna
        table.columns[0].width = Inches(0.4)  # #
        table.columns[1].width = Inches(2.0)  # Original → Corregido
        table.columns[2].width = Inches(3.5)  # Motivo
        table.columns[3].width = Inches(1.5)  # Contexto
        table.columns[4].width = Inches(0.6)  # Línea

        # Llenar tabla
        for i, e in enumerate(entries_list, start=1):
            row_cells = table.add_row().cells

            # Número
            row_cells[0].text = str(i)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Original → Corregido
            original = str(_safe_get(e, "original", ""))
            corrected = str(_safe_get(e, "corrected", ""))
            p = row_cells[1].paragraphs[0]
            p.clear()
            run_original = p.add_run(original)
            run_original.font.color.rgb = RGBColor(192, 0, 0)  # Rojo
            run_original.font.bold = True
            p.add_run(" → ")
            run_corrected = p.add_run(corrected)
            run_corrected.font.color.rgb = RGBColor(0, 176, 80)  # Verde
            run_corrected.font.bold = True

            # Motivo
            row_cells[2].text = str(_safe_get(e, "reason", ""))

            # Contexto
            context = str(_safe_get(e, "context", ""))
            row_cells[3].text = context if len(context) < 50 else context[:47] + "..."

            # Línea
            row_cells[4].text = str(_safe_get(e, "line", ""))
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Formatear fuente de la fila
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        doc.save(path)
        return
    # Fallback: write a minimal paragraphs DOCX via write_paragraphs
    paras: list[str] = []
    paras.append("Correcciones")
    if source_filename:
        paras.append(f"Documento: {source_filename}")
    paras.append(f"Total correcciones: {len(entries_list)}")
    paras.append("")
    for i, e in enumerate(entries_list, start=1):
        paras.append(
            f"{i}. Línea {_safe_get(e,'line','')} | TokenID {_safe_get(e,'token_id','')} | Chunk {_safe_get(e,'chunk_index','')}"
        )
        paras.append(f"Original: {_safe_get(e,'original','')}")
        paras.append(f"Corregido: {_safe_get(e,'corrected','')}")
        paras.append(f"Motivo: {_safe_get(e,'reason','')}")
        ctx = str(_safe_get(e, 'context', ''))
        if ctx:
            paras.append(f"Contexto: {ctx}")
        sent = str(_safe_get(e, 'sentence', ''))
        if sent:
            paras.append(f"Frase: {sent}")
        paras.append("")
    write_paragraphs(paras, path)
