from __future__ import annotations

import zipfile
from xml.etree import ElementTree as ET

try:  # optional dependency
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - optional
    Document = None  # type: ignore


def read_paragraphs(path: str) -> list[str]:
    if Document is not None and path.lower().endswith(".docx"):
        try:
            doc = Document(path)  # type: ignore
            return [p.text for p in doc.paragraphs]
        except Exception:
            pass
    if path.lower().endswith(".txt"):
        with open(path, encoding="utf-8") as f:
            return [line.rstrip("\n") for line in f]
    # Fallback: minimal DOCX parse via zip + XML
    return _read_docx_zip(path)


def write_paragraphs(paragraphs: list[str], path: str) -> None:
    if Document is not None and path.lower().endswith(".docx"):
        try:
            doc = Document()  # type: ignore
            for text in paragraphs:
                doc.add_paragraph(text)
            doc.save(path)
            return
        except Exception:
            pass
    if path.lower().endswith(".txt"):
        with open(path, "w", encoding="utf-8") as f:
            for p in paragraphs:
                f.write(p + "\n")
        return
    _write_minimal_docx(paragraphs, path)


def _read_docx_zip(path: str) -> list[str]:
    texts: list[str] = []
    with zipfile.ZipFile(path) as zf:
        with zf.open("word/document.xml") as f:
            xml = f.read()
    # Namespaces used in docx
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    root = ET.fromstring(xml)
    for p in root.findall(".//w:p", ns):
        parts: list[str] = []
        for t in p.findall(".//w:t", ns):
            parts.append(t.text or "")
        texts.append("".join(parts))
    return texts


def _write_minimal_docx(paragraphs: list[str], path: str) -> None:
    # Build a minimal docx package with only document.xml and content types
    document_xml = _build_document_xml(paragraphs)
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("word/document.xml", document_xml)


def _build_document_xml(paragraphs: list[str]) -> str:
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ET.register_namespace("w", ns)
    root = ET.Element(f"{{{ns}}}document")
    body = ET.SubElement(root, f"{{{ns}}}body")
    for ptext in paragraphs:
        p = ET.SubElement(body, f"{{{ns}}}p")
        r = ET.SubElement(p, f"{{{ns}}}r")
        t = ET.SubElement(r, f"{{{ns}}}t")
        t.text = ptext
    return ET.tostring(root, encoding="utf-8", xml_declaration=True).decode("utf-8")


def write_docx_preserving_runs(input_path: str, paragraphs: list[str], output_path: str) -> None:
    """Rewrite document.xml text while preserving run structure and formatting.

    It copies the original .docx package and only replaces the text content in w:t nodes
    to match the provided paragraphs. If the number of paragraphs differs, extra paragraphs
    are left unchanged.
    """
    with zipfile.ZipFile(input_path, "r") as zf:
        xml = zf.read("word/document.xml")
        infolist = zf.infolist()
        files = {
            i.filename: zf.read(i.filename) for i in infolist if i.filename != "word/document.xml"
        }

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root = ET.fromstring(xml)
    paras = root.findall(".//w:body/w:p", ns)

    for idx, p in enumerate(paras):
        if idx >= len(paragraphs):
            break
        new_text = paragraphs[idx]
        t_nodes = list(p.findall(".//w:t", ns))
        if not t_nodes:
            r = ET.SubElement(p, f"{{{ns['w']}}}r")
            t = ET.SubElement(r, f"{{{ns['w']}}}t")
            t_nodes = [t]
        prev_lengths = [len(t.text or "") for t in t_nodes]
        total_prev = sum(prev_lengths)
        pos = 0
        if total_prev == 0:
            t_nodes[0].text = new_text
            for t in t_nodes[1:]:
                t.text = ""
        else:
            for i, t in enumerate(t_nodes):
                length = prev_lengths[i]
                if i == len(t_nodes) - 1:
                    chunk = new_text[pos:]
                else:
                    chunk = new_text[pos : pos + length]
                t.text = chunk
                pos += len(chunk)

    new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(output_path, "w") as zf:
        for name, content in files.items():
            zf.writestr(name, content)
        zf.writestr("word/document.xml", new_xml)
